"""
Resume Parser Service

Extracts structured information from PDF/DOCX/plain-text resumes.

Two strategies, in priority order:
  1. LLM extraction  - uses Gemini structured output for accurate parsing.
  2. Regex fallback  - deterministic, offline parser used when AI is
                       unavailable (no key / expired key / SDK missing) or
                       when the AI call fails. The fallback is intentionally
                       robust so the product works fully offline.
"""

import re
import sys
from pathlib import Path
from typing import List, Optional

from pydantic import BaseModel, Field

# Add project root to path
PROJECT_ROOT = Path(__file__).parent.parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from app.utils.logger import get_logger
from app.services.ai_client import get_ai_client
from app.models.resume import (
    Resume, PersonalInfo, Education, Experience, Project, Skills
)

logger = get_logger(__name__)


# --------------------------------------------------------------------------- #
# Pydantic schemas used only for LLM structured output.
# These mirror the dataclasses in app.models.resume but are kept separate so
# the public data model stays decoupled from the AI layer.
# --------------------------------------------------------------------------- #
class _PersonalInfoSchema(BaseModel):
    name: str = ""
    email: str = ""
    phone: str = ""
    linkedin: str = ""
    github: str = ""
    location: str = ""
    portfolio: str = ""


class _EducationSchema(BaseModel):
    institution: str = ""
    degree: str = ""
    field_of_study: str = ""
    start_date: str = ""
    end_date: str = ""
    gpa: Optional[float] = None
    achievements: List[str] = Field(default_factory=list)


class _ExperienceSchema(BaseModel):
    company: str = ""
    title: str = ""
    location: str = ""
    start_date: str = ""
    end_date: str = ""
    is_current: bool = False
    responsibilities: List[str] = Field(default_factory=list)
    technologies: List[str] = Field(default_factory=list)


class _ProjectSchema(BaseModel):
    name: str = ""
    description: str = ""
    technologies: List[str] = Field(default_factory=list)
    url: str = ""
    highlights: List[str] = Field(default_factory=list)


class _SkillsSchema(BaseModel):
    technical: List[str] = Field(default_factory=list)
    soft: List[str] = Field(default_factory=list)
    tools: List[str] = Field(default_factory=list)
    languages: List[str] = Field(default_factory=list)
    frameworks: List[str] = Field(default_factory=list)


class _ResumeSchema(BaseModel):
    personal_info: _PersonalInfoSchema = Field(default_factory=_PersonalInfoSchema)
    education: List[_EducationSchema] = Field(default_factory=list)
    experience: List[_ExperienceSchema] = Field(default_factory=list)
    skills: _SkillsSchema = Field(default_factory=_SkillsSchema)
    projects: List[_ProjectSchema] = Field(default_factory=list)
    certifications: List[str] = Field(default_factory=list)
    summary: str = ""


class ResumeParser:
    """Parser to extract structured information from resumes."""

    # Section headers. Each value is matched against a *header-like* line only.
    SECTION_PATTERNS = {
        'education': r'(?i)\b(education|academic|qualifications?)\b',
        'experience': r'(?i)\b(experience|employment|work\s*history|professional\s*experience)\b',
        'skills': r'(?i)\b(skills|technical\s*skills|competenc\w*|technologies)\b',
        'projects': r'(?i)\b(projects?|portfolio|work\s*samples)\b',
        'certifications': r'(?i)\b(certifications?|certificates?|credentials?|licen[sc]es?)\b',
        'summary': r'(?i)\b(summary|objective|profile|about)\b',
    }

    EMAIL_PATTERN = r'[\w.+-]+@[\w-]+\.[\w.-]+'
    PHONE_PATTERN = r'(?:\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    LINKEDIN_PATTERN = r'linkedin\.com/in/[\w-]+'
    GITHUB_PATTERN = r'github\.com/[\w-]+'
    URL_PATTERN = r'https?://[^\s)]+'

    def __init__(self, use_ai: bool = True):
        logger.info("Initializing ResumeParser")
        self.use_ai = use_ai
        self.ai = get_ai_client() if use_ai else None
        self._check_dependencies()

    def _check_dependencies(self):
        """Log availability of optional file-parsing libraries."""
        try:
            import fitz  # noqa: F401  (PyMuPDF)
            logger.debug("PyMuPDF available for PDF parsing")
        except ImportError:
            logger.warning("PyMuPDF not installed. PDF parsing will be unavailable.")

        try:
            import docx  # noqa: F401
            logger.debug("python-docx available for DOCX parsing")
        except ImportError:
            logger.warning("python-docx not installed. DOCX parsing will be unavailable.")

    # ------------------------------------------------------------------ #
    # Public API
    # ------------------------------------------------------------------ #
    def parse(self, file_path: str) -> Resume:
        """Parse a resume file (PDF or DOCX)."""
        file_path = Path(file_path)
        logger.info(f"Parsing resume: {file_path.name}")

        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"Resume file not found: {file_path}")

        extension = file_path.suffix.lower()
        if extension == '.pdf':
            raw_text = self._extract_pdf_text(file_path)
        elif extension == '.docx':
            raw_text = self._extract_docx_text(file_path)
        else:
            logger.error(f"Unsupported file format: {extension}")
            raise ValueError(f"Unsupported file format: {extension}")

        if not raw_text.strip():
            raise ValueError("No text could be extracted from the resume file")

        logger.debug(f"Extracted {len(raw_text)} characters from resume")
        resume = self._parse_text(raw_text)
        resume.raw_text = raw_text
        logger.info(f"Successfully parsed resume for: {resume.personal_info.name or 'Unknown'}")
        return resume

    def parse_text(self, text: str) -> Resume:
        """Parse resume from raw text."""
        logger.info("Parsing resume from text input")
        if not text or not text.strip():
            raise ValueError("Resume text is empty")
        resume = self._parse_text(text)
        resume.raw_text = text
        return resume

    # ------------------------------------------------------------------ #
    # File text extraction
    # ------------------------------------------------------------------ #
    def _extract_pdf_text(self, file_path: Path) -> str:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.error("PyMuPDF not installed. Run: pip install pymupdf")
            raise ImportError("PyMuPDF required for PDF parsing. Install with: pip install pymupdf")

        try:
            logger.debug(f"Extracting text from PDF: {file_path}")
            doc = fitz.open(file_path)
            try:
                text_parts = [page.get_text() for page in doc]
            finally:
                doc.close()
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise

    def _extract_docx_text(self, file_path: Path) -> str:
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise ImportError("python-docx required for DOCX parsing. Install with: pip install python-docx")

        try:
            logger.debug(f"Extracting text from DOCX: {file_path}")
            doc = Document(file_path)
            text_parts = [para.text for para in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise

    # ------------------------------------------------------------------ #
    # Orchestration: LLM first, regex fallback
    # ------------------------------------------------------------------ #
    def _parse_text(self, text: str) -> Resume:
        if self.use_ai and self.ai is not None and self.ai.available:
            resume = self._llm_extract(text)
            if resume is not None:
                logger.info("Resume parsed via AI extraction")
                return resume
            logger.info("AI extraction unavailable/failed; using regex fallback")
        return self._regex_extract(text)

    def _llm_extract(self, text: str) -> Optional[Resume]:
        """Extract a structured resume using Gemini. Returns None on failure."""
        prompt = f"""You are a precise resume parser. Extract the resume below into the
provided JSON schema. Rules:
- Copy information verbatim; do NOT invent or embellish.
- Leave a field as an empty string / empty list if it is not present.
- For experience, set is_current=true if the role is ongoing (e.g. "Present").
- Split bullet points into individual responsibility/highlight strings.
- gpa must be a number if present, otherwise null.

Resume text:
\"\"\"
{text}
\"\"\""""

        parsed = self.ai.generate_json(prompt, schema=_ResumeSchema)
        if parsed is None:
            return None
        try:
            return self._schema_to_resume(parsed)
        except Exception as e:
            logger.warning(f"Failed to convert AI resume output: {e}")
            return None

    @staticmethod
    def _schema_to_resume(s: _ResumeSchema) -> Resume:
        """Convert the Pydantic extraction schema into the public dataclass model."""
        pi = s.personal_info
        resume = Resume(
            personal_info=PersonalInfo(
                name=pi.name, email=pi.email, phone=pi.phone,
                linkedin=pi.linkedin, github=pi.github,
                location=pi.location, portfolio=pi.portfolio,
            ),
            education=[
                Education(
                    institution=e.institution, degree=e.degree,
                    field_of_study=e.field_of_study,
                    start_date=e.start_date or None, end_date=e.end_date or None,
                    gpa=e.gpa, achievements=list(e.achievements),
                )
                for e in s.education
            ],
            experience=[
                Experience(
                    company=x.company, title=x.title, location=x.location,
                    start_date=x.start_date or None, end_date=x.end_date or None,
                    is_current=x.is_current,
                    responsibilities=list(x.responsibilities),
                    technologies=list(x.technologies),
                )
                for x in s.experience
            ],
            skills=Skills(
                technical=list(s.skills.technical), soft=list(s.skills.soft),
                tools=list(s.skills.tools), languages=list(s.skills.languages),
                frameworks=list(s.skills.frameworks),
            ),
            projects=[
                Project(
                    name=p.name, description=p.description,
                    technologies=list(p.technologies), url=p.url,
                    highlights=list(p.highlights),
                )
                for p in s.projects
            ],
            certifications=list(s.certifications),
            summary=s.summary,
        )
        return resume

    # ------------------------------------------------------------------ #
    # Regex extraction (offline fallback)
    # ------------------------------------------------------------------ #
    def _regex_extract(self, text: str) -> Resume:
        resume = Resume()
        resume.personal_info = self._extract_personal_info(text)
        sections = self._split_into_sections(text)

        if sections.get('education'):
            resume.education = self._extract_education(sections['education'])
        if sections.get('experience'):
            resume.experience = self._extract_experience(sections['experience'])
        if sections.get('skills'):
            resume.skills = self._extract_skills(sections['skills'])
        if sections.get('projects'):
            resume.projects = self._extract_projects(sections['projects'])
        if sections.get('certifications'):
            resume.certifications = self._extract_certifications(sections['certifications'])
        if sections.get('summary'):
            resume.summary = sections['summary'].strip()

        return resume

    def _extract_personal_info(self, text: str) -> PersonalInfo:
        info = PersonalInfo()

        email_match = re.search(self.EMAIL_PATTERN, text)
        if email_match:
            info.email = email_match.group()

        phone_match = re.search(self.PHONE_PATTERN, text)
        if phone_match:
            info.phone = phone_match.group()

        linkedin_match = re.search(self.LINKEDIN_PATTERN, text, re.IGNORECASE)
        if linkedin_match:
            info.linkedin = f"https://{linkedin_match.group()}"

        github_match = re.search(self.GITHUB_PATTERN, text, re.IGNORECASE)
        if github_match:
            info.github = f"https://{github_match.group()}"

        # Name: first plausible line near the top.
        for line in text.strip().split('\n')[:5]:
            line = line.strip()
            if (
                line and len(line) < 50
                and '@' not in line and 'http' not in line.lower()
                and not re.search(self.PHONE_PATTERN, line)
            ):
                info.name = line
                break

        return info

    def _match_section_header(self, line: str) -> Optional[str]:
        """Return a section key if ``line`` is a standalone section header.

        A header line is short, has few words, and carries no inline content
        after a colon (so 'Soft Skills: Java, Go' is treated as *content*,
        not as a new header).
        """
        if not line:
            return None

        if ':' in line:
            head, rest = line.split(':', 1)
            if rest.strip():  # has content after the colon -> not a pure header
                return None
        else:
            head = line

        head = head.strip()
        if not head or len(head) > 40 or len(head.split()) > 4:
            return None

        for name, pattern in self.SECTION_PATTERNS.items():
            if re.search(pattern, head):
                return name
        return None

    def _split_into_sections(self, text: str) -> dict:
        """Split resume text into named sections.

        Content lines before the first recognised header are ignored (they are
        usually the contact block). Repeated headers append rather than
        overwrite, so sub-labelled blocks are never lost.
        """
        sections: dict = {}
        current = None
        buf: List[str] = []

        for raw in text.split('\n'):
            header = self._match_section_header(raw.strip())
            if header:
                if current and buf:
                    existing = sections.get(current, '')
                    sections[current] = (existing + '\n' + '\n'.join(buf)).strip()
                current = header
                buf = []
            elif current:
                buf.append(raw)

        if current and buf:
            existing = sections.get(current, '')
            sections[current] = (existing + '\n' + '\n'.join(buf)).strip()

        return sections

    def _extract_education(self, text: str) -> List[Education]:
        education_list = []
        degree_pattern = (
            r"(?i)(Ph\.?D\.?|Doctorate|Master(?:'s)?(?:\s+of\s+\w+)?|M\.?S\.?|M\.?A\.?|MBA|"
            r"Bachelor(?:'s)?(?:\s+of\s+\w+)?|B\.?S\.?|B\.?A\.?|B\.?Tech|M\.?Tech|Associate)"
        )

        for entry in re.split(r'\n\s*\n', text):
            entry = entry.strip()
            if not entry:
                continue

            edu = Education()
            lines = entry.split('\n')

            degree_match = re.search(degree_pattern, entry)
            if degree_match:
                edu.degree = degree_match.group().strip()

            # Prefer the field after "in" ("Bachelor of Science in Computer Science"
            # -> "Computer Science"); fall back to the field after "of".
            field_match = re.search(r'(?i)\bin[^\S\n]+([A-Z][\w&]*(?:[^\S\n]+[A-Z][\w&]*){0,3})', entry)
            if not field_match:
                field_match = re.search(
                    r'(?i)\bof[^\S\n]+([A-Z][\w&]*(?:[^\S\n]+[A-Z][\w&]*){0,3}[^\S\n]*'
                    r'(?:Science|Engineering|Technology|Arts))', entry)
            if field_match:
                edu.field_of_study = field_match.group(1).strip()

            gpa_match = re.search(r'(?i)(?:GPA|CGPA)[:\s]*([\d.]+)', entry)
            if gpa_match:
                try:
                    edu.gpa = float(gpa_match.group(1))
                except ValueError:
                    pass

            year_match = re.search(r'(20\d{2}|19\d{2})', entry)
            if year_match:
                edu.end_date = year_match.group()

            # Institution: the first line that is not the degree line.
            for line in lines:
                line = line.strip()
                if line and (not degree_match or degree_match.group() not in line):
                    edu.institution = re.split(r'\s*[|,]\s*', line)[0].strip()
                    break
            if not edu.institution and lines:
                edu.institution = lines[0].strip()

            if edu.institution or edu.degree:
                education_list.append(edu)

        return education_list

    def _extract_experience(self, text: str) -> List[Experience]:
        experience_list = []

        for entry in re.split(r'\n\s*\n', text):
            entry = entry.strip()
            if not entry:
                continue

            exp = Experience()
            lines = [ln.strip() for ln in entry.split('\n') if ln.strip()]
            if not lines:
                continue

            # First non-bullet line: title. Second: company (| location | dates).
            non_bullets = [ln for ln in lines if not re.match(r'^[•\-\*◦▪]', ln)]
            if non_bullets:
                exp.title = non_bullets[0]
                if len(non_bullets) > 1:
                    parts = re.split(r'\s*[|,]\s*', non_bullets[1])
                    exp.company = parts[0].strip()
                    if len(parts) > 1:
                        exp.location = parts[1].strip()

            date_match = re.search(
                r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s*\d{4}'
                r'|\b(?:19|20)\d{2}\b)', entry, re.IGNORECASE)
            if date_match:
                exp.start_date = date_match.group()

            if re.search(r'(?i)\b(present|current|now|ongoing)\b', entry):
                exp.is_current = True

            exp.responsibilities = [
                b.strip() for b in re.findall(r'[•\-\*◦▪]\s*(.+)', entry) if b.strip()
            ]

            if exp.title or exp.company:
                experience_list.append(exp)

        return experience_list

    def _extract_skills(self, text: str) -> Skills:
        skills = Skills()

        # Map labelled lines (e.g. "Tools: Git, Docker") into categories.
        label_map = {
            'language': 'languages', 'programming': 'languages',
            'framework': 'frameworks', 'librar': 'frameworks',
            'tool': 'tools', 'technolog': 'tools', 'platform': 'tools',
            'soft': 'soft', 'interpersonal': 'soft',
        }

        consumed_lines = set()
        for line in text.split('\n'):
            if ':' not in line:
                continue
            label, _, items_str = line.partition(':')
            label_lower = label.strip().lower()
            target = 'technical'
            for key, attr in label_map.items():
                if key in label_lower:
                    target = attr
                    break
            items = [i.strip() for i in re.split(r'[,;•|]', items_str) if i.strip()]
            getattr(skills, target).extend(items)
            consumed_lines.add(line.strip())

        # Unlabelled lines: treat comma/bullet separated tokens as technical.
        for line in text.split('\n'):
            line = line.strip()
            if not line or line in consumed_lines:
                continue
            for item in re.split(r'[,;•|]', line):
                item = item.strip()
                if 2 < len(item) < 40:
                    skills.technical.append(item)

        # De-duplicate while preserving order, case-insensitively.
        for attr in ('technical', 'soft', 'tools', 'languages', 'frameworks'):
            seen = set()
            deduped = []
            for s in getattr(skills, attr):
                if s.lower() not in seen:
                    seen.add(s.lower())
                    deduped.append(s)
            setattr(skills, attr, deduped)

        return skills

    def _extract_projects(self, text: str) -> List[Project]:
        projects = []

        for entry in re.split(r'\n\s*\n', text):
            entry = entry.strip()
            if not entry:
                continue

            proj = Project()
            lines = [ln.strip() for ln in entry.split('\n') if ln.strip()]
            if not lines:
                continue

            non_bullets = [ln for ln in lines if not re.match(r'^[•\-\*◦▪]', ln)]
            if non_bullets:
                proj.name = non_bullets[0]
                if len(non_bullets) > 1:
                    proj.description = ' '.join(non_bullets[1:])

            proj.highlights = [
                b.strip() for b in re.findall(r'[•\-\*◦▪]\s*(.+)', entry) if b.strip()
            ]

            tech_match = re.search(
                r'(?i)(?:technologies?|tech\s*stack|built\s+with|tools?)[:\s]+(.+)', entry)
            if tech_match:
                proj.technologies = [
                    t.strip() for t in re.split(r'[,;]', tech_match.group(1)) if t.strip()
                ]
                # Remove the tech line from highlights to avoid duplication.
                proj.highlights = [
                    h for h in proj.highlights
                    if not re.match(r'(?i)\s*(technologies?|tech\s*stack|built\s+with)', h)
                ]

            url_match = re.search(self.URL_PATTERN, entry)
            if url_match:
                proj.url = url_match.group()

            if proj.name:
                projects.append(proj)

        return projects

    def _extract_certifications(self, text: str) -> List[str]:
        certifications = []
        for line in text.split('\n'):
            line = re.sub(r'^[•\-\*◦▪]\s*', '', line.strip())
            if line and len(line) > 3:
                certifications.append(line)
        return certifications
